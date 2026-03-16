
import io
import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from bs4 import BeautifulSoup, NavigableString
from lxml import etree

def set_word_compatibility_mode(document):
    """Define o modo de compatibilidade do documento para Word 2019/365.
    Remove a mensagem de 'Modo de Compatibilidade' ao abrir o documento.
    """
    # Word 2019/365 usa compatibilityMode = 15
    # Word 2016 = 15, Word 2013 = 15, Word 2010 = 14, Word 2007 = 12
    settings = document.settings.element
    
    # Encontrar ou criar o elemento compat
    compat = settings.find(qn('w:compat'))
    if compat is None:
        compat = OxmlElement('w:compat')
        settings.append(compat)
    
    # Remover compatibilityMode existente se houver
    for child in compat.findall(qn('w:compatSetting')):
        if child.get(qn('w:name')) == 'compatibilityMode':
            compat.remove(child)
    
    # Adicionar compatibilityMode para Word 2019 (valor 15)
    compat_setting = OxmlElement('w:compatSetting')
    compat_setting.set(qn('w:name'), 'compatibilityMode')
    compat_setting.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
    compat_setting.set(qn('w:val'), '15')  # Word 2019/365
    compat.append(compat_setting)

def set_run_font_and_color(run, font_size=None, bold=None, italic=None):
    """Define a fonte e a cor de uma execução."""
    font = run.font
    font.name = 'Open Sans'
    font.color.rgb = RGBColor(0, 0, 0)
    if font_size:
        font.size = font_size
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic


def set_table_borders(table):
    """Aplica bordas a todas as células da tabela (equivalente ao estilo 'Table Grid')."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # tamanho da borda
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # cor preta
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_table(document, headers, data, header_bold=True, column_widths=None):
    """Adiciona uma tabela ao documento."""
    if not data:
        return
    table = document.add_table(rows=1, cols=len(headers))
    
    # Aplica bordas manualmente (compatível com templates sem estilos)
    set_table_borders(table)

    if column_widths:
        for i, width in enumerate(column_widths):
            if i < len(table.columns):
                table.columns[i].width = width

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header
        
        # Centraliza o conteúdo do cabeçalho
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                if header_bold:
                    run.font.bold = True
                set_run_font_and_color(run, font_size=Pt(7))

        # Adiciona fundo cinza ao cabeçalho
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')  # Cinza médio (25% branco)
        cell._tc.get_or_add_tcPr().append(shading_elm)


    for item in data:
        row_cells = table.add_row().cells
        if isinstance(item, dict):
            for i, key in enumerate(item.keys()):
                cell = row_cells[i]
                cell.text = str(item.get(key, ''))
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font_and_color(run, font_size=Pt(7))
        else:
            cell = row_cells[0]
            cell.text = str(item)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font_and_color(run, font_size=Pt(7))

def add_formatted_text_to_doc(p, text_content, is_markdown=False):
    """Analisa uma string HTML ou Markdown e adiciona conteúdo formatado a um parágrafo do docx."""
    if is_markdown:
        html_content = markdown.markdown(text_content)
    else:
        html_content = text_content

    soup = BeautifulSoup(f"<div>{html_content}</div>", 'html.parser')
    for element in soup.div.contents:
        if isinstance(element, NavigableString):
            run = p.add_run(str(element))
            set_run_font_and_color(run)
        elif element.name in ['b', 'strong']:
            run = p.add_run(element.get_text())
            set_run_font_and_color(run, bold=True)
        elif element.name in ['i', 'em']:
            run = p.add_run(element.get_text())
            set_run_font_and_color(run, italic=True)
        else:
            run = p.add_run(element.get_text())
            set_run_font_and_color(run)

def format_estrategia_nome_for_display(strategy_key: str) -> str:
    """Formata o nome da estratégia para exibição."""
    if not strategy_key:
        return ""
    return strategy_key.replace("-", " ").title()

def _add_subtopicos(document, subtopicos: list, prefixo: str = "", nivel: int = 1):
    """Adiciona subtópicos de conhecimentos com numeração hierárquica ao documento."""
    indent = "    " * nivel  # Indentação baseada no nível
    for idx, subtopico in enumerate(subtopicos, start=1):
        descricao = subtopico.get('descricao', '')
        if descricao:
            numero = f"{prefixo}{idx}"
            p = document.add_paragraph()
            run = p.add_run(f"{indent}{numero}. {descricao}")
            set_run_font_and_color(run)
            # Processar sub-subtópicos recursivamente
            sub_subtopicos = subtopico.get('subtopicos', [])
            if sub_subtopicos:
                _add_subtopicos(document, sub_subtopicos, f"{numero}.", nivel + 1)

def _generate_docx_from_ai_plan(document, plan_content: dict):
    """Gera o conteúdo do DOCX para um plano gerado por IA."""
    # Informações do Curso
    info_curso = plan_content.get('plano_de_ensino', {}).get('informacoes_curso', {})
    document.add_heading('1. Informações do Curso', level=1)
    table = document.add_table(rows=0, cols=2)
    set_table_borders(table)  # Aplica bordas (compatível com templates)
    table.columns[0].width = Inches(2) # Define a largura da primeira coluna
    table.columns[1].width = Inches(4) # Define a largura da segunda coluna
    for key, value in info_curso.items():
        row_cells = table.add_row().cells
        # Célula do Título (Negrito)
        # Célula do Título (Negrito e com fundo)
        cell_key = row_cells[0]
        p_key = cell_key.paragraphs[0]
        run_key = p_key.add_run(f"{key.replace('_', ' ').title()}:")
        run_key.bold = True
        set_run_font_and_color(run_key, font_size=Pt(11))
        
        # Adiciona fundo cinza à célula do título
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')  # Cinza médio
        cell_key._tc.get_or_add_tcPr().append(shading_elm)

        # Célula do Valor
        p_value = row_cells[1].paragraphs[0]
        run_value = p_value.add_run(str(value))
        set_run_font_and_color(run_value, font_size=Pt(11))

    # Situações de Aprendizagem
    document.add_section()
    document.sections[-1].orientation = WD_ORIENT.PORTRAIT
    situacoes = plan_content.get('plano_de_ensino', {}).get('situacoes_aprendizagem', [])
    document.add_heading('2. Situações de Aprendizagem', level=1)

    for i, sa in enumerate(situacoes):
        document.add_heading(f"2.{i+1}. {sa.get('titulo', 'Situação de Aprendizagem')}", level=2)

        # Capacidades a Desenvolver
        document.add_heading('Capacidades a Desenvolver', level=3)
        capacidades = sa.get('capacidades', {})
        
        # Capacidades Básicas/Técnicas
        capacidades_basicas = capacidades.get('basicas', []) or capacidades.get('tecnicas', [])
        if capacidades_basicas:
            p = document.add_paragraph()
            run = p.add_run("Capacidades Básicas/Técnicas:")
            set_run_font_and_color(run, bold=True)
            for idx, cap in enumerate(capacidades_basicas, start=1):
                p = document.add_paragraph()
                run = p.add_run(f"    {idx}. {cap}")
                set_run_font_and_color(run)
        
        # Capacidades Socioemocionais
        capacidades_socioemocionais = capacidades.get('socioemocionais', [])
        if capacidades_socioemocionais:
            p = document.add_paragraph()
            run = p.add_run("Capacidades Socioemocionais:")
            set_run_font_and_color(run, bold=True)
            for idx, cap in enumerate(capacidades_socioemocionais, start=1):
                p = document.add_paragraph()
                run = p.add_run(f"    {idx}. {cap}")
                set_run_font_and_color(run)

        # Conhecimentos (com hierarquia completa numerada)
        document.add_heading('Conhecimentos', level=3)
        conhecimentos = sa.get('conhecimentos', [])
        if conhecimentos:
            for idx, conhecimento in enumerate(conhecimentos, start=1):
                topico = conhecimento.get('topico', '')
                if topico:
                    # Tópico principal (negrito e numerado)
                    p = document.add_paragraph()
                    run = p.add_run(f"{idx}. {topico}")
                    set_run_font_and_color(run, bold=True)
                    
                    # Subtópicos recursivamente com numeração hierárquica
                    _add_subtopicos(document, conhecimento.get('subtopicos', []), prefixo=f"{idx}.", nivel=1)

        # Estratégia de Aprendizagem
        estrategia = sa.get('estrategia_aprendizagem', {})
        # Note: 'estrategia' key (slug) might be lost in Markdown->JSON conversion.
        # We try to use 'tipo' from the extracted strategy content which usually contains the name.
        estrategia_tipo = estrategia.get('tipo')
        
        titulo_estrategia = 'Estratégia de Aprendizagem'
        if estrategia_tipo:
             # Clean up potential extra spacing or punctuation from LLM extraction if needed
             nome_estrategia = estrategia_tipo.strip().replace(":", "") 
             titulo_estrategia = f"{titulo_estrategia} - {nome_estrategia}"

        document.add_section()
        document.sections[-1].orientation = WD_ORIENT.PORTRAIT     
        document.add_heading(titulo_estrategia, level=3)
        detalhes_estrategia = estrategia.get('detalhes', {})
        for key, value in detalhes_estrategia.items():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(f"{key.replace('_', ' ').title()}: ")
            set_run_font_and_color(run, bold=True)
            add_formatted_text_to_doc(p, value, is_markdown=True)

        document.add_section()
        document.sections[-1].orientation = WD_ORIENT.PORTRAIT
        # Critérios de Avaliação Dicotômicos
        document.add_heading('Critérios de Avaliação - Dicotômicos', level=3)
        
        # Instrumento de Registro - Cabeçalho
        p = document.add_paragraph()
        run = p.add_run("Instrumento de Registro")
        set_run_font_and_color(run, bold=True)
        p = document.add_paragraph()
        run = p.add_run("Nome do aluno: ______________________    Turma: ______________________")
        set_run_font_and_color(run)
        
        criterios_dicotomicos = sa.get('criterios_avaliacao', {}).get('dicotomicos', [])
        headers_dicotomicos = ["Capacidade", "Critérios de Avaliação", "Autoavaliação", "Avaliação"]
        data_dicotomicos = []
        for c in criterios_dicotomicos:
            criterios_list = c.get('criterios', [])
            # Primeira linha com a capacidade e primeiro critério
            if criterios_list:
                data_dicotomicos.append({
                    "capacidade": c.get('capacidade', ''),
                    "criterios": criterios_list[0] if len(criterios_list) > 0 else '',
                    "autoavaliacao": "",
                    "avaliacao": ""
                })
                # Linhas adicionais para critérios restantes (sem repetir capacidade)
                for crit in criterios_list[1:]:
                    data_dicotomicos.append({
                        "capacidade": "",
                        "criterios": crit,
                        "autoavaliacao": "",
                        "avaliacao": ""
                    })
        add_table(document, headers_dicotomicos, data_dicotomicos)
        
        # Legenda
        p = document.add_paragraph()
        run = p.add_run("Legenda: S = Atingiu / N = Não Atingiu")
        set_run_font_and_color(run, italic=True)

        document.add_section()
        document.sections[-1].orientation = WD_ORIENT.PORTRAIT
        # Critérios de Avaliação Graduais
        document.add_heading('Critérios de Avaliação - Graduais', level=3)
        
        # Instrumento de Registro - Cabeçalho
        p = document.add_paragraph()
        run = p.add_run("Instrumento de Registro")
        set_run_font_and_color(run, bold=True)
        p = document.add_paragraph()
        run = p.add_run("Nome do aluno: ______________________    Turma: ______________________")
        set_run_font_and_color(run)
        
        criterios_graduais = sa.get('criterios_avaliacao', {}).get('graduais', [])
        
        # Verificar se existe o campo 'criterios' (lista não vazia) em algum item para decidir as colunas
        has_criterios_column = any(c.get('criterios') for c in criterios_graduais)
        
        if has_criterios_column:
            headers_graduais = ["Capacidade", "Critérios", "Nível 1", "Nível 2", "Nível 3", "Nível 4", "Nível Alcançado"]
        else:
            headers_graduais = ["Capacidade", "Nível 1", "Nível 2", "Nível 3", "Nível 4", "Nível Alcançado"]

        data_graduais = []
        for c in criterios_graduais:
            niveis = c.get('niveis', {})
            row_data = {}
            # Ordem de inserção DEVE corresponder à ordem dos headers para add_table funcionar corretamente
            row_data["capacidade"] = c.get('capacidade')
            
            if has_criterios_column:
                criterios_list = c.get('criterios', [])
                if isinstance(criterios_list, list):
                    row_data["criterios"] = "\n".join([str(x) for x in criterios_list])
                else:
                    row_data["criterios"] = str(criterios_list) if criterios_list else ""

            row_data["nivel_1"] = niveis.get('nivel_1')
            row_data["nivel_2"] = niveis.get('nivel_2')
            row_data["nivel_3"] = niveis.get('nivel_3')
            row_data["nivel_4"] = niveis.get('nivel_4')
            row_data["nivel_alcancado"] = ""
            
            data_graduais.append(row_data)

        add_table(document, headers_graduais, data_graduais)
        
        # Legenda
        p = document.add_paragraph()
        run = p.add_run("Legenda:")
        set_run_font_and_color(run, bold=True)
        
        p = document.add_paragraph()
        run = p.add_run("Nível 1: Desempenho autônomo – apresenta desempenho esperado da competência com autonomia, sem intervenções do docente.")
        set_run_font_and_color(run, italic=True)
        
        p = document.add_paragraph()
        run = p.add_run("Nível 2: Desempenho parcialmente autônomo – apresenta desempenho esperado da competência, com intervenções pontuais do docente.")
        set_run_font_and_color(run, italic=True)
        
        p = document.add_paragraph()
        run = p.add_run("Nível 3: Desempenho apoiado – ainda não apresenta desempenho esperado da competência, exigindo intervenções constantes do docente.")
        set_run_font_and_color(run, italic=True)
        
        p = document.add_paragraph()
        run = p.add_run("Nível 4: Desempenho não satisfatório – ainda não apresenta desempenho esperado da competência, mesmo com intervenções constantes do docente.")
        set_run_font_and_color(run, italic=True)

        # --- Seção Paisagem para Plano de Aula ---
        section_landscape = document.add_section()
        section_landscape.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section_landscape.page_height, section_landscape.page_width
        section_landscape.page_width = new_width
        section_landscape.page_height = new_height

        # Configurar margens específicas para esta seção
        section_landscape.top_margin = Inches(0.5)
        section_landscape.bottom_margin = Inches(0.5)
        section_landscape.left_margin = Inches(0.5)
        section_landscape.right_margin = Inches(0.5)

        document.add_heading('Plano de Aula', level=3)
        plano_aula = sa.get('plano_de_aula', [])
        headers_plano_aula = ["Horas/Aulas/Data", "Capacidades", "Conhecimentos", "Estratégias", "Recursos/Ambientes", "Critérios de Avaliação", "Instrumento de Avaliação", "Referências"]
        column_widths_plano_aula = [Inches(1.5), Inches(1), Inches(1), Inches(1), Inches(1.5), Inches(1.5), Inches(1.7), Inches(1.7)]
        add_table(document, headers_plano_aula, plano_aula, column_widths=column_widths_plano_aula)

        # --- Retorna para Seção Retrato ---
        section_portrait = document.add_section()
        section_portrait.orientation = WD_ORIENT.PORTRAIT
        new_width, new_height = section_portrait.page_height, section_portrait.page_width
        section_portrait.page_width = new_width
        section_portrait.page_height = new_height

        # Configurar margens específicas para esta seção
        section_portrait.top_margin = Inches(0.5)
        section_portrait.bottom_margin = Inches(0.5)
        section_portrait.left_margin = Inches(1)
        section_portrait.right_margin = Inches(1)

def _generate_docx_from_manual_plan(document, plan_content: dict):
    """Gera o conteúdo do DOCX para um plano criado manualmente."""
    info_gerais = plan_content.get('informacoes_gerais', {})
    document.add_heading('1. Informações Gerais', level=1)
    table = document.add_table(rows=0, cols=2)
    set_table_borders(table)  # Aplica bordas (compatível com templates)
    table.columns[0].width = Inches(1.7) # Define a largura da primeira coluna
    table.columns[1].width = Inches(5.0) # Define a largura da segunda coluna
    info_map = {
        "curso": "Curso", "turma": "Turma", "unidade_curricular": "Unidade Curricular",
        "professor": "Professor", "escola": "Escola", "departamento_regional": "Departamento Regional"
    }
    for key, label in info_map.items():
        value = info_gerais.get(key, '')
        row_cells = table.add_row().cells
        # Célula do Título (Negrito)
        cell = row_cells[0]
        p_key = cell.paragraphs[0]
        run_key = p_key.add_run(f"{label}:")
        run_key.bold = True
        set_run_font_and_color(run_key, font_size=Pt(11))
        
        # Adiciona fundo cinza ao cabeçalho
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')  # Cinza médio (25% branco)
        cell._tc.get_or_add_tcPr().append(shading_elm)

        # Célula do Valor
        p_value = row_cells[1].paragraphs[0]
        run_value = p_value.add_run(str(value))
        set_run_font_and_color(run_value, font_size=Pt(11))

    situacoes = plan_content.get('situacoes_aprendizagem', [])
    document.add_heading('2. Situações de Aprendizagem', level=1)

    for i, sa in enumerate(situacoes):
        document.add_heading(f"2.{i+1}. {sa.get('tema', 'Situação de Aprendizagem')}", level=2)
        
        # Capacidades e Conhecimentos vêm ANTES da Situação de Aprendizagem
        p = document.add_paragraph()
        add_formatted_text_to_doc(p, "Capacidades Técnicas: " + ", ".join(sa.get('capacidades_tecnicas', [])))
        p = document.add_paragraph()
        add_formatted_text_to_doc(p, "Capacidades Socioemocionais: " + ", ".join(sa.get('capacidades_socioemocionais', [])))
        p = document.add_paragraph()
        add_formatted_text_to_doc(p, "Conhecimentos: " + ", ".join(sa.get('conhecimentos', [])))
        
        # Situação de Aprendizagem (Desafio) vem depois
        desafio_html = sa.get('desafio', '')
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run("Situação de Aprendizagem: ")
        set_run_font_and_color(run, bold=True)
        add_formatted_text_to_doc(p, desafio_html, is_markdown=True)

        document.add_heading('Critérios de Avaliação', level=3)
        for crit in sa.get('criterios', []):
            p = document.add_paragraph()
            run = p.add_run(f"Critério ({crit.get('tipo', '')}): ")
            run.bold = True
            set_run_font_and_color(run)
            set_run_font_and_color(run)
            add_formatted_text_to_doc(p, crit.get('criterio', ''), is_markdown=True)
            if crit.get('tipo') == 'gradual':
                for j in range(1, 5):
                    if crit.get(f'nivel{j}'):
                        p = document.add_paragraph(style='List Bullet')
                        add_formatted_text_to_doc(p, f"Nível {j}: {crit.get(f'nivel{j}')}", is_markdown=True)

        # --- Seção Paisagem para Plano de Aula ---
        section_landscape = document.add_section()
        section_landscape.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section_landscape.page_height, section_landscape.page_width
        section_landscape.page_width = new_width
        section_landscape.page_height = new_height

        # Configurar margens específicas para esta seção
        section_landscape.top_margin = Inches(0.5)
        section_landscape.bottom_margin = Inches(0.5)
        section_landscape.left_margin = Inches(0.5)
        section_landscape.right_margin = Inches(0.5)

        document.add_heading('Plano de Aula', level=3)
        plano_aula_data = sa.get('plano_aula', [])
        headers_plano_aula = ["Horas/Aulas/Data", "Capacidades", "Conhecimentos", "Estratégias", "Recursos/Ambientes", "Critérios de Avaliação", "Instrumento de Avaliação", "Referências"]
        column_widths_plano_aula = [Inches(1.5), Inches(1), Inches(1), Inches(1), Inches(1.5), Inches(1.5), Inches(1.7), Inches(1.7)]

        data_plano_aula_formatted = []
        for aula in plano_aula_data:
            data_plano_aula_formatted.append({
                "Horas/Aulas/Data": f"{aula.get('data', '').split('T')[0]} {aula.get('hora_inicio', '')}-{aula.get('hora_fim', '')}",
                "Capacidades": ", ".join(aula.get('capacidades', [])),
                "Conhecimentos": ", ".join(aula.get('conhecimentos', [])),
                "Estratégias": ", ".join(aula.get('estrategias', [])),
                "Recursos/Ambientes": ", ".join(aula.get('recursos_ambientes', [])),
                "Critérios de Avaliação": ", ".join(aula.get('criterios', [])), # Assuming 'criterios' for manual plan
                "Instrumento de Avaliação": ", ".join(aula.get('instrumentos', [])), # Assuming 'instrumentos' for manual plan
                "Referências": ", ".join(aula.get('referencias', [])),
            })
        add_table(document, headers_plano_aula, data_plano_aula_formatted, column_widths=column_widths_plano_aula)
        
        # --- Retorna para Seção Retrato ---
        section_portrait = document.add_section()
        section_portrait.orientation = WD_ORIENT.PORTRAIT
        new_width, new_height = section_portrait.page_height, section_portrait.page_width
        section_portrait.page_width = new_width
        section_portrait.page_height = new_height

def generate_docx(plan_content: dict, departamento_regional: str | None):
    """Gera um documento DOCX a partir do conteúdo do plano de ensino."""
    # document = Document()
    document = Document('src/template/template.docx')
    
    # Define modo de compatibilidade para Word 2019/365
    # set_word_compatibility_mode(document)

    # Escolhe o logo com base no departamento regional fornecido
    if departamento_regional == 'SP':
        logo_path = 'src/imgs/logo_senai_sp.png'
    else:
        logo_path = 'src/imgs/logo_senai.png'

    header = document.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = header_paragraph.add_run()
    logo_run.add_picture(logo_path, width=Inches(2.0))
    header.add_paragraph() # Adiciona uma linha em branco após o logo

    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run('Plano de Ensino')
    title_run.bold = True
    set_run_font_and_color(title_run, font_size=Pt(24))
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if 'plano_de_ensino' in plan_content:
        _generate_docx_from_ai_plan(document, plan_content)
    elif 'informacoes_gerais' in plan_content:
        _generate_docx_from_manual_plan(document, plan_content)
    else:
        raise ValueError("Formato de JSON do plano desconhecido.")

    for p in document.paragraphs:
        if p.style.name.startswith('Heading'):
            for r in p.runs:
                set_run_font_and_color(r)

    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io
