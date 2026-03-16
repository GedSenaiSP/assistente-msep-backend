"""
Agente LangGraph para geração de Cadernos de Estudo.

Gera material didático personalizado para Situações de Aprendizagem do SENAI.
Estrutura com SA na introdução, capítulos técnicos e páginas de anotações.
"""
import os
import json
import logging
import time
import uuid
from io import BytesIO
from typing import Dict, List, Any, TypedDict, Optional, Generator
from pathlib import Path

from langgraph.graph import StateGraph, END
from langchain_openai import AzureChatOpenAI

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
import re

# Configuração de logs
logger = logging.getLogger(__name__)


class DidacticResourceState(TypedDict, total=False):
    """Estado do agente de geração de Caderno de Estudo."""
    # Inputs
    sa_tema: str  # Tema/desafio da SA
    sa_descricao: str  # Descrição completa da SA
    sa_capacidades_tecnicas: List[str]
    sa_capacidades_socioemocionais: List[str]
    sa_conhecimentos: List[Dict[str, Any]]
    sa_estrategia: str  # situacao-problema, projeto, etc.
    curso_nome: str
    uc_nome: str
    area_tecnologica: str
    num_chapters: int
    
    # Estado interno
    title: str
    outline: List[Dict[str, Any]]
    chapters: Dict[int, Dict[str, str]]
    current_chapter: int
    status: str
    feedback: str
    
    # Output
    docx_bytes: bytes
    error: str


def get_llm() -> AzureChatOpenAI:
    """Retorna instância do AzureChatOpenAI para geração de conteúdo."""
    return AzureChatOpenAI(
        azure_deployment=os.getenv("MODEL_ID", "gpt-4o"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        temperature=0.7,
        max_tokens=65535
    )


def safe_json_parse(response_text: str, fallback: Any) -> Any:
    """Tenta decodificar JSON e retorna um fallback em caso de erro."""
    text = response_text.strip()
    if text.startswith("```json"):
        text = text[7:]
        if text.endswith("```"):
            text = text[:-3]
    elif text.startswith("```"):
        text = text[3:]
        if text.endswith("```"):
            text = text[:-3]
    
    # Tentar encontrar JSON no texto
    start_bracket = text.find('[')
    start_brace = text.find('{')
    
    if start_bracket == -1 and start_brace == -1:
        logger.error(f"Nenhum JSON encontrado no texto: {text[:100]}...")
        return fallback
    
    try:
        return json.loads(text.strip())
    except json.JSONDecodeError:
        logger.error(f"Erro ao decodificar JSON: {text[:100]}... Usando fallback.")
        return fallback


async def generate_title(state: DidacticResourceState) -> Dict[str, Any]:
    """Gera título para o Caderno de Estudo baseado na SA."""
    logger.info("Gerando título do Caderno de Estudo...")
    llm = get_llm()
    
    prompt = f"""
    Você é um especialista em educação profissional do SENAI. 
    Crie um título formal e técnico para um CADERNO DE ESTUDO que será usado como material de apoio para uma Situação de Aprendizagem.
    
    Informações da Situação de Aprendizagem:
    - Tema/Desafio: {state.get('sa_tema', 'Não especificado')}
    - Curso: {state.get('curso_nome', 'Não especificado')}
    - Unidade Curricular: {state.get('uc_nome', 'Não especificada')}
    - Estratégia: {state.get('sa_estrategia', 'Não especificada')}
    
    Capacidades Técnicas a desenvolver:
    {chr(10).join(['- ' + cap for cap in state.get('sa_capacidades_tecnicas', [])])}
    
    O título deve:
    - Ser técnico e objetivo
    - Refletir o conteúdo técnico abordado
    - Ter no máximo 60 caracteres
    
    Responda SOMENTE em formato JSON com a chave "title", sem texto adicional.
    Exemplo: {{"title": "Fundamentos de Automação Industrial"}}
    """
    
    response = await llm.ainvoke(prompt)
    info = safe_json_parse(response.content, {"title": f"Caderno de Estudo: {state.get('sa_tema', 'Recurso')}"})
    
    logger.info(f"Título gerado: {info.get('title')}")
    
    return {
        "title": info.get("title", f"Caderno de Estudo: {state.get('sa_tema', 'Recurso')}"),
        "status": "title_generated"
    }


async def create_outline(state: DidacticResourceState) -> Dict[str, Any]:
    """Cria o sumário do Caderno de Estudo."""
    logger.info("Criando sumário do Caderno de Estudo...")
    llm = get_llm()
    
    num_chapters = state.get("num_chapters", 3)
    
    # Extrair conhecimentos
    conhecimentos_text = ""
    if state.get("sa_conhecimentos"):
        for conhecimento in state["sa_conhecimentos"]:
            if isinstance(conhecimento, dict):
                topico = conhecimento.get("topico", "")
                if topico:
                    conhecimentos_text += f"- {topico}\n"
    
    prompt = f"""
    Você é um especialista em educação profissional do SENAI, criando um Caderno de Estudo para uma Situação de Aprendizagem.
    
    Informações:
    - Título: {state.get('title')}
    - Tema/Desafio: {state.get('sa_tema')}
    - Curso: {state.get('curso_nome')}
    - Unidade Curricular: {state.get('uc_nome')}
    - Estratégia de Aprendizagem: {state.get('sa_estrategia')}
    
    Capacidades Técnicas a desenvolver:
    {chr(10).join(['- ' + cap for cap in state.get('sa_capacidades_tecnicas', [])])}
    
    Conhecimentos técnicos identificados:
    {conhecimentos_text if conhecimentos_text else 'Extrair dos capacidades técnicas'}
    
    ESTRUTURA DO CADERNO:
    O Caderno de Estudo deve ter {num_chapters} capítulos de conhecimentos técnicos que embasem a execução da SA.
    
    - NÃO inclua capítulo de "Introdução" (a SA será apresentada separadamente)
    - NÃO inclua capítulo de "Conclusão" ou "Considerações Finais"
    - Cada capítulo deve focar em um CONHECIMENTO TÉCNICO específico
    - Os capítulos devem seguir progressão didática (do fundamento ao avançado)
    
    Responda SOMENTE em formato JSON com uma lista de objetos:
    [{{"chapter_number": 1, "chapter_title": "Nome do Conhecimento Técnico", "chapter_description": "Descrição breve do conteúdo"}}]
    """
    
    response = await llm.ainvoke(prompt)
    outline_data = safe_json_parse(response.content, [
        {"chapter_number": 1, "chapter_title": "Fundamentos Técnicos", 
         "chapter_description": f"Introdução aos conceitos fundamentais relacionados a {state.get('sa_tema')}."}
    ])
    
    # Garantir número correto de capítulos
    if len(outline_data) < num_chapters:
        for i in range(len(outline_data) + 1, num_chapters + 1):
            outline_data.append({
                "chapter_number": i,
                "chapter_title": f"Conhecimentos Técnicos - Parte {i}",
                "chapter_description": f"Continuação do desenvolvimento técnico."
            })
    
    chapters = {
        item["chapter_number"]: {
            "title": item["chapter_title"], 
            "description": item["chapter_description"],
            "content": ""
        } 
        for item in outline_data[:num_chapters]
    }
    
    logger.info(f"Sumário criado com {len(chapters)} capítulos.")
    
    return {
        "outline": outline_data[:num_chapters],
        "chapters": chapters,
        "current_chapter": 1,
        "status": "outline_created"
    }


async def write_chapter(state: DidacticResourceState) -> Dict[str, Any]:
    """Escreve o conteúdo de um capítulo do Caderno de Estudo."""
    current = state["current_chapter"]
    
    if current > len(state["chapters"]):
        logger.info("Todos os capítulos foram escritos.")
        return {"status": "all_chapters_written"}
    
    chapter_info = state["chapters"][current]
    logger.info(f"Escrevendo Capítulo {current}: {chapter_info['title']}...")
    
    llm = get_llm()
    
    # Contexto do capítulo anterior
    prev_context = ""
    if current > 1 and state["chapters"].get(current-1, {}).get("content"):
        prev_chapter = state["chapters"][current-1]
        prev_context = f"""
        O capítulo anterior ({prev_chapter['title']}) abordou:
        {prev_chapter['content'][:300]}...
        
        Continue a progressão didática a partir desse conhecimento.
        """
    
    prompt = f"""
    Você é um especialista em educação profissional do SENAI, escrevendo um Caderno de Estudo intitulado "{state['title']}".
    
    Contexto da Situação de Aprendizagem:
    - Tema/Desafio: {state.get('sa_tema')}
    - Curso: {state.get('curso_nome')}
    - Unidade Curricular: {state.get('uc_nome')}
    - Estratégia de Aprendizagem: {state.get('sa_estrategia')}
    
    Escreva o Capítulo {current}: "{chapter_info['title']}".
    
    Descrição do capítulo: {chapter_info['description']}
    
    {prev_context}
    
    DIRETRIZES IMPORTANTES:
    1. Escreva conteúdo TÉCNICO de qualidade, como em um livro didático profissional
    2. O texto deve embasar a execução da Situação de Aprendizagem
    3. Aproximadamente 2000-3000 palavras
    4. Inclua conceitos técnicos detalhados com explicações claras
    5. Adicione exemplos práticos do mundo do trabalho
    6. Inclua dicas, alertas de segurança quando aplicável
    
    REGRAS OBRIGATÓRIAS:
    - NÃO inclua NENHUM texto introdutório ou conversacional
    - NÃO escreva frases como "Com certeza", "Olá", "Vamos lá", "Dando continuidade", etc.
    - NÃO repita o título do capítulo no início
    - NÃO inclua saudações ou apresentações
    - Comece DIRETAMENTE com a primeira seção [H2] do conteúdo técnico
    
    FORMATAÇÃO - Use APENAS estas marcações para estruturar:
    - Use [H2] para títulos de seções principais
    - Use [H3] para subtítulos
    - Use [H4] para subsubtítulos
    - Use [BOLD] e [/BOLD] para texto em negrito
    - Use [LISTA] antes de cada item de lista com bullet
    - Use [NUMLISTA] antes de cada item de lista numerada
    
    Exemplo de formatação:
    [H2]Nome da Seção
    
    Parágrafo de texto normal com [BOLD]termos importantes[/BOLD] destacados.
    
    [H3]Subtítulo
    
    [LISTA] Primeiro item da lista
    [LISTA] Segundo item da lista
    
    NÃO use markdown (#, ##, **, -, etc.). Use APENAS as marcações indicadas acima.
    """
    
    response = await llm.ainvoke(prompt)
    
    # Garantir que o conteúdo seja string
    content = response.content
    if isinstance(content, list):
        content = "\n".join([str(c) for c in content])
    
    updated_chapters = state["chapters"].copy()
    updated_chapters[current]["content"] = content
    
    logger.info(f"Capítulo {current} concluído.")
    
    next_chapter = current + 1
    new_status = "chapter_written" if next_chapter <= len(state["chapters"]) else "all_chapters_written"
    
    return {
        "chapters": updated_chapters,
        "current_chapter": next_chapter,
        "status": new_status
    }


def set_font_all_scripts(font, font_name: str):
    """Define a fonte para todos os scripts (Latin, EastAsia, ComplexScript).
    
    Isso é necessário porque os estilos de Heading do Word usam fontes do tema,
    e apenas definir font.name não sobrescreve essas fontes do tema.
    """
    font.name = font_name
    # Acessar o elemento XML do run properties
    element = font.element
    rPr = element
    
    # Encontrar ou criar elemento rFonts
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    
    # Definir fonte para todos os scripts
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def setup_document_styles(doc: Document):
    """Configura estilos do documento para sumário automático."""
    styles = doc.styles
    
    # Estilo Heading 1 - Capítulos
    if 'Heading 1' in styles:
        h1 = styles['Heading 1']
        set_font_all_scripts(h1.font, 'Open Sans')
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(22, 65, 148)  # Azul escuro SENAI
    
    # Estilo Heading 2 - Seções
    if 'Heading 2' in styles:
        h2 = styles['Heading 2']
        set_font_all_scripts(h2.font, 'Open Sans')
        h2.font.size = Pt(14)
        h2.font.bold = False
        h2.font.color.rgb = RGBColor(22, 65, 148)
    
    # Estilo Heading 3 - Subseções
    if 'Heading 3' in styles:
        h3 = styles['Heading 3']
        set_font_all_scripts(h3.font, 'Open Sans')
        h3.font.size = Pt(14)
        h3.font.bold = False
        h3.font.italic = True
        h3.font.color.rgb = RGBColor(22, 65, 148)

    # Estilo Heading 4 - Subsubseções
    if 'Heading 4' in styles:
        h4 = styles['Heading 4']
        set_font_all_scripts(h4.font, 'Open Sans')
        h4.font.size = Pt(12)
        h4.font.bold = False
        h4.font.italic = False
        h4.font.color.rgb = RGBColor(22, 65, 148)
    
    # Estilo Normal
    if 'Normal' in styles:
        normal = styles['Normal']
        set_font_all_scripts(normal.font, 'Open Sans')
        normal.font.size = Pt(12)
        normal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def add_toc_placeholder(doc: Document):
    """Adiciona campo de Sumário automático do Word."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Criar campo TOC
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-4" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    # Instrução para atualizar
    note = doc.add_paragraph()
    note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = note.add_run("(Clique com o botão direito e selecione 'Atualizar campo' para gerar o sumário)")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_notes_page(doc: Document, chapter_title: str):
    """Adiciona uma página de anotações estilo caderno."""
    doc.add_page_break()
    
    # Título da página
    heading = doc.add_heading("Anotações", level=2)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Subtítulo com referência ao capítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run(f"Espaço para anotações sobre: {chapter_title}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Espaço
    
    # Criar linhas de anotação
    for i in range(22):  # ~22 linhas por página
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(28)  # Espaçamento entre linhas
        # Adicionar linha horizontal
        run = p.add_run("_" * 96)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(10)


def process_content_to_docx(text: str, doc: Document):
    """Processa o conteúdo com marcações personalizadas para DOCX."""
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Heading 2 - Seções principais
        if line.startswith('[H2]'):
            title_text = line[4:].strip()
            doc.add_heading(title_text, level=2)
            continue
        
        # Heading 3 - Subseções
        if line.startswith('[H3]'):
            title_text = line[4:].strip()
            doc.add_heading(title_text, level=3)
            continue
        
        # Heading 4 - Subsubseções
        if line.startswith('[H4]'):
            title_text = line[4:].strip()
            doc.add_heading(title_text, level=4)
            continue
        
        # Lista com bullet
        if line.startswith('[LISTA]'):
            item_text = line[7:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text_new(item_text, p)
            continue
        
        # Lista numerada
        if line.startswith('[NUMLISTA]'):
            item_text = line[10:].strip()
            p = doc.add_paragraph(style='List Number')
            add_formatted_text_new(item_text, p)
            continue
        
        # Fallback para markdown (caso o LLM ainda use)
        if line.startswith('##'):
            level = min(line.count('#', 0, 6), 6)
            title_text = line.lstrip('#').strip()
            doc.add_heading(title_text, level=min(level, 3))
            continue
        
        if re.match(r'^[-*]\s+', line):
            item_text = re.sub(r'^[-*]\s+', '', line).strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text_new(item_text, p)
            continue
        
        if re.match(r'^\d+\.\s+', line):
            item_text = re.sub(r'^\d+\.\s+', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            add_formatted_text_new(item_text, p)
            continue
        
        # Parágrafo normal
        paragraph = doc.add_paragraph()
        add_formatted_text_new(line, paragraph)


def add_formatted_text_new(text: str, paragraph):
    """Adiciona texto formatado com marcações [BOLD] ao parágrafo."""
    # Processar [BOLD]...[/BOLD]
    pattern = r'\[BOLD\](.*?)\[/BOLD\]'
    parts = re.split(pattern, text)
    
    # Também processar **...** como fallback
    is_bold = False
    for i, part in enumerate(parts):
        if not part:
            continue
        
        # Partes ímpares são o conteúdo de [BOLD]
        if i % 2 == 1:
            run = paragraph.add_run(part)
            run.bold = True
        else:
            # Processar ** markdown como fallback
            md_parts = re.split(r'(\*\*.*?\*\*)', part)
            for md_part in md_parts:
                if md_part.startswith('**') and md_part.endswith('**'):
                    run = paragraph.add_run(md_part[2:-2])
                    run.bold = True
                else:
                    run = paragraph.add_run(md_part)


async def export_docx(state: DidacticResourceState) -> Dict[str, Any]:
    """Exporta o Caderno de Estudo para DOCX com formatação profissional."""
    logger.info("Exportando Caderno de Estudo para DOCX...")
    
    doc = Document()
    
    # Configurar estilos
    setup_document_styles(doc)
    
    # === CAPA ===
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Título do Caderno
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("CADERNO DE ESTUDO")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(22, 65, 148)
    
    doc.add_paragraph()
    
    # Título do Material
    main_title = doc.add_paragraph()
    main_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    mt_run = main_title.add_run(state.get("title", "Material Didático"))
    mt_run.font.size = Pt(20)
    mt_run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Informações do Curso
    info_para = doc.add_paragraph()
    info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    info_para.add_run(f"Curso: ").bold = True
    info_para.add_run(state.get("curso_nome", "") + "\n")
    info_para.add_run(f"Unidade Curricular: ").bold = True
    info_para.add_run(state.get("uc_nome", "") + "\n")
    info_para.add_run(f"Estratégia de Aprendizagem: ").bold = True
    info_para.add_run(state.get("sa_estrategia", "").replace("-", " ").title())
    
    doc.add_page_break()
    
    # === SUMÁRIO ===
    doc.add_heading("Sumário", level=1)
    add_toc_placeholder(doc)
    doc.add_page_break()
    
    # === INTRODUÇÃO - APRESENTAÇÃO DA SA ===
    doc.add_heading("Apresentação da Situação de Aprendizagem", level=1)
    
    # Tema/Desafio
    doc.add_heading("Contextualização", level=2)
    intro_para = doc.add_paragraph()
    intro_para.add_run("Este Caderno de Estudo foi desenvolvido para apoiar a execução da Situação de Aprendizagem proposta, fornecendo os conhecimentos técnicos necessários para o desenvolvimento das capacidades previstas.")
    
    doc.add_paragraph()
    
    tema_para = doc.add_paragraph()
    tema_para.add_run("Tema/Desafio: ").bold = True
    tema_para.add_run(state.get("sa_tema", "Não especificado"))
    
    # Capacidades Técnicas
    doc.add_heading("Capacidades Técnicas a Desenvolver", level=2)
    for cap in state.get("sa_capacidades_tecnicas", []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(cap)
    
    # Capacidades Socioemocionais
    if state.get("sa_capacidades_socioemocionais"):
        doc.add_heading("Capacidades Socioemocionais", level=2)
        for cap in state.get("sa_capacidades_socioemocionais", []):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(cap)
    
    # Como usar o caderno
    doc.add_heading("Como Utilizar Este Caderno", level=2)
    uso_para = doc.add_paragraph()
    uso_para.add_run("Este material foi estruturado para auxiliar você no desenvolvimento das capacidades técnicas necessárias para resolver o desafio proposto. Ao final de cada capítulo, há um espaço dedicado para suas anotações pessoais, dúvidas e observações importantes.")
    
    doc.add_paragraph()
    dica_para = doc.add_paragraph()
    dica_para.add_run("Dica: ").bold = True
    dica_para.add_run("Use as páginas de anotações para registrar suas reflexões, conectar o conteúdo com a prática e anotar pontos importantes discutidos em sala de aula.")
    
    doc.add_page_break()
    
    # === CAPÍTULOS ===
    for chapter_num, chapter_data in sorted(state["chapters"].items()):
        # Título do capítulo
        doc.add_heading(f"Capítulo {chapter_num}: {chapter_data['title']}", level=1)
        
        # Conteúdo do capítulo
        process_content_to_docx(chapter_data["content"], doc)
        
        # Página de anotações após cada capítulo
        add_notes_page(doc, chapter_data['title'])
        
        # Quebra de página (exceto no último)
        if chapter_num < len(state["chapters"]):
            doc.add_page_break()
    
    # === REFERÊNCIAS (placeholder) ===
    # doc.add_page_break()
    # doc.add_heading("Referências Bibliográficas", level=1)
    # ref_note = doc.add_paragraph()
    # ref_note.add_run("As referências bibliográficas utilizadas na elaboração deste material seguem as normas da ABNT e estão disponíveis para consulta na biblioteca do SENAI.")
    # ref_note.paragraph_format.space_after = Pt(12)
    
    # Salvar em bytes
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()
    
    logger.info(f"Caderno de Estudo exportado com sucesso ({len(docx_bytes)} bytes).")
    
    return {
        "docx_bytes": docx_bytes,
        "status": "exported"
    }


def router(state: DidacticResourceState) -> str:
    """Decide o próximo estado do workflow."""
    status_map = {
        "start": "generate_title",
        "title_generated": "create_outline",
        "outline_created": "write_chapter",
        "chapter_written": "write_chapter",
        "all_chapters_written": "export_docx",
        "exported": END
    }
    next_state = status_map.get(state.get("status", "start"), END)
    logger.debug(f"Transição de estado: {state.get('status')} -> {next_state}")
    return next_state


def create_didactic_resource_agent() -> StateGraph:
    """Cria o agente de geração de Cadernos de Estudo."""
    logger.info("Criando agente de geração de Cadernos de Estudo...")
    
    workflow = StateGraph(DidacticResourceState)
    
    workflow.add_node("generate_title", generate_title)
    workflow.add_node("create_outline", create_outline)
    workflow.add_node("write_chapter", write_chapter)
    workflow.add_node("export_docx", export_docx)
    
    workflow.set_entry_point("generate_title")
    
    workflow.add_conditional_edges("generate_title", router)
    workflow.add_conditional_edges("create_outline", router)
    workflow.add_conditional_edges("write_chapter", router)
    workflow.add_conditional_edges("export_docx", router)
    
    return workflow.compile()


async def run_didactic_resource_agent(
    sa_tema: str,
    sa_capacidades_tecnicas: List[str],
    sa_capacidades_socioemocionais: List[str],
    sa_estrategia: str,
    curso_nome: str,
    uc_nome: str,
    num_chapters: int = 3,
    area_tecnologica: str = "",
    sa_descricao: str = "",
    sa_conhecimentos: List[Dict[str, Any]] = None,
    on_progress: Optional[callable] = None
) -> Dict[str, Any]:
    """
    Executa o agente de geração de Cadernos de Estudo.
    
    Args:
        sa_tema: Tema/desafio da Situação de Aprendizagem
        sa_capacidades_tecnicas: Lista de capacidades técnicas
        sa_capacidades_socioemocionais: Lista de capacidades socioemocionais
        sa_estrategia: Estratégia de aprendizagem (situacao-problema, projeto, etc.)
        curso_nome: Nome do curso
        uc_nome: Nome da unidade curricular
        num_chapters: Número de capítulos (1-10)
        area_tecnologica: Área tecnológica (opcional)
        sa_descricao: Descrição completa da SA (opcional)
        sa_conhecimentos: Conhecimentos técnicos da SA (opcional)
        on_progress: Callback para atualizar progresso (progress, step)
    
    Returns:
        Dict com docx_bytes, title, status
    """
    logger.info(f"Iniciando geração de Caderno de Estudo para SA: {sa_tema}")
    
    agent = create_didactic_resource_agent()
    
    initial_state = DidacticResourceState(
        sa_tema=sa_tema,
        sa_capacidades_tecnicas=sa_capacidades_tecnicas,
        sa_capacidades_socioemocionais=sa_capacidades_socioemocionais,
        sa_estrategia=sa_estrategia,
        curso_nome=curso_nome,
        uc_nome=uc_nome,
        num_chapters=min(num_chapters, 10),
        area_tecnologica=area_tecnologica,
        sa_descricao=sa_descricao,
        sa_conhecimentos=sa_conhecimentos or [],
        status="start"
    )
    
    config = {"recursion_limit": 100}
    
    # Calcular total de etapas para progresso
    total_steps = 2 + num_chapters + 1  # título + sumário + N capítulos + export
    current_step = 0
    
    try:
        # Executar o agente
        async for output in agent.astream(initial_state, config=config):
            node_name = list(output.keys())[0] if output else "unknown"
            node_output = output.get(node_name, {})
            status = node_output.get("status", "unknown")
            
            # Atualizar progresso
            if on_progress:
                if status == "title_generated":
                    current_step = 1
                    on_progress(int((current_step / total_steps) * 100), "Gerando título...")
                elif status == "outline_created":
                    current_step = 2
                    on_progress(int((current_step / total_steps) * 100), "Criando estrutura...")
                elif status == "chapter_written":
                    written_chapter = node_output.get("current_chapter", 1) - 1
                    if written_chapter > 0:
                        current_step = 2 + written_chapter
                        on_progress(int((current_step / total_steps) * 100), f"Escrevendo capítulo {written_chapter}...")
                elif status == "exported":
                    current_step = total_steps
                    on_progress(100, "Gerando documento...")
            
            # Atualizar estado
            initial_state.update(node_output)
        
        logger.info("Geração de Caderno de Estudo concluída com sucesso!")
        
        return {
            "docx_bytes": initial_state.get("docx_bytes"),
            "title": initial_state.get("title"),
            "status": "completed"
        }
        
    except Exception as e:
        logger.error(f"Erro durante a geração do Caderno de Estudo: {e}", exc_info=True)
        return {
            "error": str(e),
            "status": "failed"
        }
