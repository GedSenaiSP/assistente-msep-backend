"""
Agente LangGraph para geração de listas de exercícios.
"""
import os
import json
import logging
import re
from io import BytesIO
from typing import Dict, List, Any, TypedDict

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from langgraph.graph import StateGraph, END
from langchain_openai import AzureChatOpenAI

logger = logging.getLogger(__name__)


class ExercisesState(TypedDict, total=False):
    sa_tema: str
    sa_capacidades_tecnicas: List[str]
    sa_conhecimentos: List[Dict[str, Any]]
    quantities: Dict[str, int]
    generated_questions: Dict[str, List[Any]]
    docx_bytes: bytes
    error: str
    status: str


def get_llm():
    return AzureChatOpenAI(
        azure_deployment=os.getenv("MODEL_ID", "gpt-4o"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        temperature=0.7,
        max_tokens=65535,
        model_kwargs={"response_format": {"type": "json_object"}},
    )


def safe_json_parse(text: str, fallback: Any) -> Any:
    text = text.strip()
    if text.startswith("```json"):
        text = text[7:]
    if text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    try:
        return json.loads(text.strip())
    except:
        return fallback


def add_formatted_text(paragraph, text: str, base_font_size: int = 12):
    """Adiciona texto ao parágrafo processando marcações [BOLD]...[/BOLD].
    
    Também converte markdown como fallback:
    - [BOLD]texto[/BOLD] -> negrito
    - **texto** -> negrito (fallback)
    """
    if not text:
        return
    
    # Primeiro converter markdown para marcações personalizadas
    # **bold** -> [BOLD]
    text = re.sub(r'\*\*(.*?)\*\*', r'[BOLD]\1[/BOLD]', text)
    # __bold__ -> [BOLD] (mas NÃO captura múltiplos underscores consecutivos como lacunas)
    # Só captura quando há texto que não é underscore entre os __
    text = re.sub(r'__([^_]+)__', r'[BOLD]\1[/BOLD]', text)
    # Remover outras formatações markdown que não queremos
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove *italic*
    # Remove _italic_ mas NÃO remove múltiplos underscores consecutivos (lacunas)
    # Só remove quando há exatamente um _ de cada lado com texto no meio
    text = re.sub(r'(?<![_])_([^_]+)_(?![_])', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)     # Remove `code`
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)  # Remove headers
    
    # Processar [BOLD]...[/BOLD]
    pattern = r'\[BOLD\](.*?)\[/BOLD\]'
    parts = re.split(pattern, text)
    
    for i, part in enumerate(parts):
        if not part:
            continue
        
        run = paragraph.add_run(part)
        run.font.size = Pt(base_font_size)
        set_font_all_scripts(run.font, 'Open Sans')
        
        # Partes ímpares são o conteúdo de [BOLD]
        if i % 2 == 1:
            run.bold = True


def clean_text(text: str) -> str:
    """Remove marcações e retorna texto limpo."""
    if not text:
        return ""
    text = re.sub(r'\[BOLD\](.*?)\[/BOLD\]', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # __bold__ - só remove quando há texto que não é underscore entre os __
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove _italic_ mas NÃO remove múltiplos underscores consecutivos (lacunas)
    text = re.sub(r'(?<![_])_([^_]+)_(?![_])', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    return text.strip()


def set_font_all_scripts(font, font_name: str):
    """Define a fonte para todos os scripts (Latin, EastAsia, ComplexScript)."""
    font.name = font_name
    element = font.element
    rPr = element
    
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def setup_document_styles(doc: Document):
    """Configura estilos do documento (mesmo do caderno de estudos)."""
    styles = doc.styles
    
    # Estilo Heading 1
    if 'Heading 1' in styles:
        h1 = styles['Heading 1']
        set_font_all_scripts(h1.font, 'Open Sans')
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(22, 65, 148)  # Azul escuro SENAI
    
    # Estilo Heading 2
    if 'Heading 2' in styles:
        h2 = styles['Heading 2']
        set_font_all_scripts(h2.font, 'Open Sans')
        h2.font.size = Pt(14)
        h2.font.bold = False
        h2.font.color.rgb = RGBColor(22, 65, 148)
    
    # Estilo Heading 3
    if 'Heading 3' in styles:
        h3 = styles['Heading 3']
        set_font_all_scripts(h3.font, 'Open Sans')
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(22, 65, 148)
    
    # Estilo Normal
    if 'Normal' in styles:
        normal = styles['Normal']
        set_font_all_scripts(normal.font, 'Open Sans')
        normal.font.size = Pt(12)
        normal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def add_question_paragraph(doc: Document, number: int, text: str):
    """Adiciona parágrafo de questão formatado com suporte a [BOLD]."""
    p = doc.add_paragraph()
    # Número em negrito
    run_num = p.add_run(f"{number}. ")
    run_num.bold = True
    run_num.font.size = Pt(12)
    set_font_all_scripts(run_num.font, 'Open Sans')
    # Texto da questão com formatação
    add_formatted_text(p, text, base_font_size=12)
    return p


def add_option_paragraph(doc: Document, option_text: str):
    """Adiciona opção de resposta formatada com suporte a [BOLD]."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    add_formatted_text(p, option_text, base_font_size=11)
    return p


async def generate_questions(state: ExercisesState) -> Dict:
    logger.info("Gerando questões...")
    llm = get_llm()
    
    conhecimentos_text = ""
    for c in state.get("sa_conhecimentos", []):
        if isinstance(c, dict) and c.get("topico"):
            conhecimentos_text += f"- {c.get('topico')}\n"
        elif isinstance(c, str):
            conhecimentos_text += f"- {c}\n"

    quantities = state.get("quantities", {})
    results = {}
    
    # 1. Múltipla Escolha
    qtd_mc = quantities.get("multiple_choice", 0)
    if qtd_mc > 0:
        prompt_mc = f"""Gere {qtd_mc} questões de MÚLTIPLA ESCOLHA sobre:
{conhecimentos_text}

REGRAS OBRIGATÓRIAS:
- Nível TÉCNICO profissional
- Cada questão deve ter exatamente 4 alternativas (A, B, C, D)
- Para destacar termos importantes, use [BOLD]termo[/BOLD]
- NÃO use markdown (**, *, _, `, #, etc.) - use apenas [BOLD]...[/BOLD]
- Alternativas devem ser completas e claras

Responda APENAS em JSON válido, sem texto adicional:
[{{"question": "Qual é a função do [BOLD]termo técnico[/BOLD] em determinado contexto?", "options": ["A) Primeira opção", "B) Segunda opção", "C) Terceira opção", "D) Quarta opção"], "correct_answer": "A) Primeira opção"}}]"""
        resp = await llm.ainvoke(prompt_mc)
        results["multiple_choice"] = safe_json_parse(resp.content, [])

    # 2. Dissertativas
    qtd_essay = quantities.get("essay", 0)
    if qtd_essay > 0:
        prompt_essay = f"""Gere {qtd_essay} questões DISSERTATIVAS sobre:
{conhecimentos_text}

REGRAS OBRIGATÓRIAS:
- Nível TÉCNICO profissional
- Para destacar termos importantes, use [BOLD]termo[/BOLD]
- NÃO use markdown (**, *, _, `, #, etc.) - use apenas [BOLD]...[/BOLD]
- Questões devem exigir análise ou explicação detalhada

Responda APENAS em JSON válido, sem texto adicional:
[{{"question": "Explique o conceito de [BOLD]termo técnico[/BOLD] e sua aplicação prática.", "expected_answer": "Tópicos esperados: primeiro ponto, segundo ponto, etc."}}]"""
        resp = await llm.ainvoke(prompt_essay)
        results["essay"] = safe_json_parse(resp.content, [])

    # 3. Completar Lacunas (com 4 opções)
    qtd_fill = quantities.get("fill_in_the_blank", 0)
    if qtd_fill > 0:
        prompt_fill = f"""Gere {qtd_fill} questões de COMPLETAR LACUNAS sobre:
{conhecimentos_text}

REGRAS OBRIGATÓRIAS:
- Nível TÉCNICO profissional
- Use '__________' (10 underlines) para indicar a lacuna
- Cada questão deve ter exatamente 4 opções de resposta (A, B, C, D)
- Apenas uma opção é correta
- Para destacar termos, use [BOLD]termo[/BOLD]
- NÃO use markdown (**, *, _, `, #, etc.) - use apenas [BOLD]...[/BOLD]

Responda APENAS em JSON válido, sem texto adicional:
[{{"sentence": "O conceito de __________ é fundamental para...", "options": ["A) termo1", "B) termo2", "C) termo3", "D) termo4"], "correct_answer": "A) termo1"}}]"""
        resp = await llm.ainvoke(prompt_fill)
        results["fill_in_the_blank"] = safe_json_parse(resp.content, [])

    # 4. Práticas
    qtd_prac = quantities.get("practical", 0)
    if qtd_prac > 0:
        prompt_prac = f"""Gere {qtd_prac} EXERCÍCIOS PRÁTICOS (Lab, Projeto ou Estudo de Caso) sobre:
{conhecimentos_text}

REGRAS OBRIGATÓRIAS:
- Nível TÉCNICO profissional
- Para destacar termos importantes, use [BOLD]termo[/BOLD]
- NÃO use markdown (**, *, _, `, #, etc.) - use apenas [BOLD]...[/BOLD]
- Descrição detalhada do que deve ser feito
- Critérios claros de avaliação

Responda APENAS em JSON válido, sem texto adicional:
[{{"title": "Título da atividade", "description": "Descrição detalhada do exercício usando [BOLD]termos importantes[/BOLD].", "criteria": "Critérios de avaliação: primeiro critério, segundo critério, etc."}}]"""
        resp = await llm.ainvoke(prompt_prac)
        results["practical"] = safe_json_parse(resp.content, [])

    # 5. Associar Colunas (Matching)
    qtd_match = quantities.get("matching", 0)
    if qtd_match > 0:
        prompt_match = f"""Gere {qtd_match} questões de ASSOCIAR COLUNAS sobre:
{conhecimentos_text}

REGRAS OBRIGATÓRIAS:
- Nível TÉCNICO profissional
- Cada questão deve ter entre 4 e 6 itens para associar
- A Coluna A contém os termos/conceitos
- A Coluna B contém as definições/descrições (embaralhadas)
- Para destacar termos, use [BOLD]termo[/BOLD]
- NÃO use markdown (**, *, _, `, #, etc.) - use apenas [BOLD]...[/BOLD]

Responda APENAS em JSON válido, sem texto adicional:
[{{"instruction": "Associe os conceitos da Coluna A com suas definições na Coluna B.", "column_a": ["1. Termo1", "2. Termo2", "3. Termo3", "4. Termo4"], "column_b": ["A) Definição do termo 2", "B) Definição do termo 4", "C) Definição do termo 1", "D) Definição do termo 3"], "answers": ["1-C", "2-A", "3-D", "4-B"]}}]"""
        resp = await llm.ainvoke(prompt_match)
        results["matching"] = safe_json_parse(resp.content, [])

    return {"generated_questions": results, "status": "questions_generated"}


def export_docx(state: ExercisesState) -> Dict:
    logger.info("Exportando DOCX...")
    try:
        doc = Document()
        
        # Configurar estilos
        setup_document_styles(doc)
        
        # Título principal
        title_para = doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_para.add_run("LISTA DE EXERCÍCIOS")
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(22, 65, 148)
        set_font_all_scripts(title_run.font, 'Open Sans')
        
        # Subtítulo com tema
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle_para.add_run(clean_text(state.get('sa_tema', '')))
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.bold = True
        set_font_all_scripts(subtitle_run.font, 'Open Sans')
        
        doc.add_paragraph()
        
        questions = state.get("generated_questions", {})
        section_number = 1
        
        # Seção: Múltipla Escolha
        if questions.get("multiple_choice"):
            doc.add_heading(f"{section_number}. Questões de Múltipla Escolha", level=1)
            section_number += 1
            
            for i, q in enumerate(questions["multiple_choice"]):
                add_question_paragraph(doc, i + 1, q.get('question', ''))
                doc.add_paragraph()
                for opt in q.get("options", []):
                    add_option_paragraph(doc, opt)
                doc.add_paragraph()

        # Seção: Completar Lacunas (agora com opções)
        if questions.get("fill_in_the_blank"):
            doc.add_heading(f"{section_number}. Complete as Lacunas", level=1)
            section_number += 1
            
            for i, q in enumerate(questions["fill_in_the_blank"]):
                add_question_paragraph(doc, i + 1, q.get('sentence', ''))
                doc.add_paragraph()
                # Adicionar opções se existirem
                for opt in q.get("options", []):
                    add_option_paragraph(doc, opt)
                doc.add_paragraph()

        # Seção: Dissertativas (sem linhas para resposta)
        if questions.get("essay"):
            doc.add_heading(f"{section_number}. Questões Dissertativas", level=1)
            section_number += 1
            
            for i, q in enumerate(questions["essay"]):
                add_question_paragraph(doc, i + 1, q.get('question', ''))
                doc.add_paragraph()

        # Seção: Atividades Práticas
        if questions.get("practical"):
            doc.add_heading(f"{section_number}. Atividades Práticas", level=1)
            section_number += 1
            
            for i, q in enumerate(questions["practical"]):
                # Título da atividade
                heading = doc.add_heading(f"Atividade {i + 1}: {clean_text(q.get('title', ''))}", level=2)
                
                # Descrição com formatação
                desc_para = doc.add_paragraph()
                add_formatted_text(desc_para, q.get('description', ''), base_font_size=12)
                
                doc.add_paragraph()

        # Seção: Associar Colunas
        if questions.get("matching"):
            doc.add_heading(f"{section_number}. Associar Colunas", level=1)
            section_number += 1
            
            for i, q in enumerate(questions["matching"]):
                # Instrução da questão
                add_question_paragraph(doc, i + 1, q.get('instruction', 'Associe os itens da Coluna A com a Coluna B.'))
                doc.add_paragraph()
                
                # Criar tabela com duas colunas SEM bordas
                column_a = q.get('column_a', [])
                column_b = q.get('column_b', [])
                max_rows = max(len(column_a), len(column_b))
                
                if max_rows > 0:
                    # Criar tabela: cabeçalho + linhas de dados
                    table = doc.add_table(rows=max_rows + 1, cols=2)
                    
                    # Remover bordas da tabela (tabela invisível)
                    tbl = table._tbl
                    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
                    tblBorders = OxmlElement('w:tblBorders')
                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'nil')
                        tblBorders.append(border)
                    tblPr.append(tblBorders)
                    if tbl.tblPr is None:
                        tbl.insert(0, tblPr)
                    
                    # Configurar largura das colunas
                    from docx.shared import Cm
                    table.columns[0].width = Cm(6)   # Coluna A
                    table.columns[1].width = Cm(10)  # Coluna B
                    
                    # Cabeçalho
                    header_cells = table.rows[0].cells
                    
                    # Coluna A header
                    header_a = header_cells[0].paragraphs[0]
                    run_a = header_a.add_run("Coluna A")
                    run_a.bold = True
                    run_a.font.size = Pt(12)
                    run_a.font.color.rgb = RGBColor(22, 65, 148)
                    set_font_all_scripts(run_a.font, 'Open Sans')
                    
                    # Coluna B header
                    header_b = header_cells[1].paragraphs[0]
                    run_b = header_b.add_run("Coluna B")
                    run_b.bold = True
                    run_b.font.size = Pt(12)
                    run_b.font.color.rgb = RGBColor(22, 65, 148)
                    set_font_all_scripts(run_b.font, 'Open Sans')
                    
                    # Preencher dados
                    for j in range(max_rows):
                        row_cells = table.rows[j + 1].cells
                        
                        # Coluna A - apenas o item numerado
                        item_a = column_a[j] if j < len(column_a) else ""
                        cell_a = row_cells[0].paragraphs[0]
                        add_formatted_text(cell_a, item_a, base_font_size=11)
                        
                        # Coluna B - parêntese + letra + definição
                        item_b = column_b[j] if j < len(column_b) else ""
                        cell_b = row_cells[1].paragraphs[0]
                        # Adicionar parêntese para resposta
                        run_paren = cell_b.add_run("(    )  ")
                        run_paren.font.size = Pt(11)
                        set_font_all_scripts(run_paren.font, 'Open Sans')
                        # Adicionar o item da coluna B
                        add_formatted_text(cell_b, item_b, base_font_size=11)
                
                doc.add_paragraph()

        # ========== GABARITO ==========
        doc.add_page_break()
        
        gabarito_title = doc.add_paragraph()
        gabarito_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        gabarito_run = gabarito_title.add_run("GABARITO")
        gabarito_run.font.size = Pt(20)
        gabarito_run.font.bold = True
        gabarito_run.font.color.rgb = RGBColor(22, 65, 148)
        set_font_all_scripts(gabarito_run.font, 'Open Sans')
        
        doc.add_paragraph()
        
        # Gabarito: Múltipla Escolha
        if questions.get("multiple_choice"):
            doc.add_heading("Múltipla Escolha", level=2)
            for i, q in enumerate(questions["multiple_choice"]):
                p = doc.add_paragraph()
                run = p.add_run(f"{i + 1}. {clean_text(q.get('correct_answer', ''))}")
                run.font.size = Pt(12)
                set_font_all_scripts(run.font, 'Open Sans')

        # Gabarito: Completar Lacunas
        if questions.get("fill_in_the_blank"):
            doc.add_heading("Complete as Lacunas", level=2)
            for i, q in enumerate(questions["fill_in_the_blank"]):
                p = doc.add_paragraph()
                # Usa correct_answer se existir (novo formato), senão usa correct_word (formato antigo)
                answer = q.get('correct_answer') or q.get('correct_word', '')
                run = p.add_run(f"{i + 1}. {clean_text(answer)}")
                run.font.size = Pt(12)
                set_font_all_scripts(run.font, 'Open Sans')

        # Gabarito: Dissertativas (expectativa de resposta)
        if questions.get("essay"):
            doc.add_heading("Questões Dissertativas (Expectativa de Resposta)", level=2)
            for i, q in enumerate(questions["essay"]):
                p = doc.add_paragraph()
                run_num = p.add_run(f"{i + 1}. ")
                run_num.bold = True
                run_num.font.size = Pt(12)
                set_font_all_scripts(run_num.font, 'Open Sans')
                
                run_text = p.add_run(clean_text(q.get('expected_answer', '')))
                run_text.font.size = Pt(12)
                set_font_all_scripts(run_text.font, 'Open Sans')

        # Gabarito: Práticas (critérios)
        if questions.get("practical"):
            doc.add_heading("Atividades Práticas (Critérios de Avaliação)", level=2)
            for i, q in enumerate(questions["practical"]):
                p = doc.add_paragraph()
                run_title = p.add_run(f"Atividade {i + 1}: ")
                run_title.bold = True
                run_title.font.size = Pt(12)
                set_font_all_scripts(run_title.font, 'Open Sans')
                
                run_criteria = p.add_run(clean_text(q.get('criteria', '')))
                run_criteria.font.size = Pt(12)
                set_font_all_scripts(run_criteria.font, 'Open Sans')

        # Gabarito: Associar Colunas
        if questions.get("matching"):
            doc.add_heading("Associar Colunas", level=2)
            for i, q in enumerate(questions["matching"]):
                p = doc.add_paragraph()
                run_num = p.add_run(f"{i + 1}. ")
                run_num.bold = True
                run_num.font.size = Pt(12)
                set_font_all_scripts(run_num.font, 'Open Sans')
                
                # Formatar as respostas (1-C, 2-A, 3-D, 4-B)
                answers = q.get('answers', [])
                answers_text = ", ".join(answers) if answers else "Não disponível"
                run_answers = p.add_run(answers_text)
                run_answers.font.size = Pt(12)
                set_font_all_scripts(run_answers.font, 'Open Sans')

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return {"docx_bytes": output.getvalue(), "status": "completed"}

    except Exception as e:
        logger.error(f"Erro export DOCX: {e}", exc_info=True)
        return {"error": str(e), "status": "failed"}


def create_graph():
    workflow = StateGraph(ExercisesState)
    workflow.add_node("generate", generate_questions)
    workflow.add_node("export", export_docx)
    workflow.set_entry_point("generate")
    workflow.add_edge("generate", "export")
    workflow.add_edge("export", END)
    return workflow.compile()


agent = create_graph()


async def generate_exercises(
    sa_tema: str,
    sa_capacidades_tecnicas: List[str],
    sa_conhecimentos: List[Dict],
    quantities: Dict[str, int]
) -> Dict:
    initial_state: ExercisesState = {
        "sa_tema": sa_tema,
        "sa_capacidades_tecnicas": sa_capacidades_tecnicas,
        "sa_conhecimentos": sa_conhecimentos,
        "quantities": quantities,
        "status": "starting"
    }
    
    try:
        result = await agent.ainvoke(initial_state)
        return {
            "docx_bytes": result.get("docx_bytes"),
            "status": "completed"
        }
    except Exception as e:
        logger.error(f"Erro: {e}")
        return {"error": str(e), "status": "failed"}
